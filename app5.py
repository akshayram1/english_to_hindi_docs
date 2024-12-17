from transformers import MarianMTModel, MarianTokenizer
from docx import Document
import torch
from tqdm import tqdm
import os

class LocalTranslator:
    def __init__(self):
        print("Loading translation model... (first time may take a few minutes)")
        self.model_name = 'Helsinki-NLP/opus-mt-en-hi'
        self.tokenizer = MarianTokenizer.from_pretrained(self.model_name)
        self.model = MarianMTModel.from_pretrained(self.model_name)
        if torch.cuda.is_available():
            self.model = self.model.to('cuda')
        print("Model loaded successfully!")
        
    def translate_text(self, text: str) -> str:
        if not text.strip():
            return text
            
        try:
            inputs = self.tokenizer(text, return_tensors="pt", padding=True, truncation=True, max_length=512)
            if torch.cuda.is_available():
                inputs = {k: v.to('cuda') for k, v in inputs.items()}
            translated = self.model.generate(**inputs)
            result = self.tokenizer.decode(translated[0], skip_special_tokens=True)
            return result
        except Exception as e:
            print(f"Translation error: {e}")
            return text

    def process_document(self, input_path: str, output_path: str) -> bool:
        try:
            doc = Document(input_path)
            
            # Count total items for progress bar
            total_items = len(doc.paragraphs) + sum(len(table.rows) * len(table.columns) for table in doc.tables)
            
            with tqdm(total=total_items, desc="Translating") as pbar:
                # Process tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                cell.text = self.translate_text(cell.text)
                            pbar.update(1)

                # Process paragraphs
                for para in doc.paragraphs:
                    if para.text.strip() and not para._element.xpath('.//w:drawing'):
                        para.text = self.translate_text(para.text)
                    pbar.update(1)

            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Document processing error: {e}")
            return False

def main():
    translator = LocalTranslator()
    input_file = "input.docx"
    output_file = "translated_hindi333333333.docx"

    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found")
        return

    if translator.process_document(input_file, output_file):
        print(f"Document translated successfully: {output_file}")
    else:
        print("Translation failed")

if __name__ == "__main__":
    main()