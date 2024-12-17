import docx
from docx.shared import Pt
from openai import OpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_openai import ChatOpenAI
import os
def translate_text_gpt4(content, llm):
    """
    Translate English text to Hindi using GPT-4 via LangChain.
    """
    messages = [
        SystemMessage(content="You are a professional Hindi translator. You are Finance Expert. You'll not translate Financial Terms and keep them as it is. Translate the following text into Hindi while retaining the meaning, context, and tone."),
        HumanMessage(content=content)
    ]
    response = llm(messages)
    return response.content.strip()
def process_docx(input_file, output_file, llm):
    """
    Process input docx file:
    1. Translate text.
    2. Preserve images in position.
    3. Export a new translated docx file.
    """
    # Load the document
    doc = docx.Document(input_file)
    # Create a new document to save the translated content
    translated_doc = docx.Document()
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraphs
            para = docx.text.paragraph.Paragraph(element, doc)
            if para.text.strip():  # Translate only non-empty text
                translated_text = translate_text_gpt4(para.text, llm)
                new_para = translated_doc.add_paragraph(translated_text)
                # Copy formatting (font size, etc.)
                for run, new_run in zip(para.runs, new_para.runs):
                    new_run.font.size = run.font.size if run.font.size else Pt(11)
        elif element.tag.endswith('tbl'):  # Tables
            table = docx.table.Table(element, doc)
            new_table = translated_doc.add_table(rows=0, cols=len(table.columns))
            for row in table.rows:
                new_row = new_table.add_row()
                for i, cell in enumerate(row.cells):
                    translated_text = translate_text_gpt4(cell.text, llm)
                    new_row.cells[i].text = translated_text
        elif element.tag.endswith('drawing'):  # Images
            # Preserve images in the same position
            translated_doc.element.body.append(element)
    # Save the translated document
    translated_doc.save(output_file)
    print(f"Translated document saved to: {output_file}")
if __name__ == "__main__":
    import os
    os.environ["OPENAI_API_KEY"] = "your open ai key"  # Set your OpenAI API key
    # Initialize the language model
    llm = ChatOpenAI(model="gpt-4o", temperature=0)
    input_file = "input.docx"   # Input English docx file
    output_file = "output_hindi_ft.docx"  # Output Hindi docx file
    if os.path.exists(input_file):
        process_docx(input_file, output_file, llm)
    else:
        print("Input file not found.")