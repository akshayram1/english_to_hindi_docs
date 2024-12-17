import requests
from docx import Document
import os
import time
from typing import Optional
from tqdm import tqdm

class AzureTranslator:
    def __init__(self):
        self.endpoint = "http://localhost:5000"
        self.headers = {
            "Content-Type": "application/json",
            "Accept": "application/json"
        }
        self.check_container()
    
    def check_container(self):
        try:
            response = requests.get(f"{self.endpoint}/")
            if response.status_code != 200:
                print("Warning: Translation container not responding")
                print("Run: docker run --name translator -d -p 5000:5000 mcr.microsoft.com/azure-cognitive-services/translator")
        except:
            print("Error: Container not running")
            
    def translate_text(self, text: str) -> str:
        if not text.strip():
            return text
            
        try:
            payload = [{"Text": text}]
            response = requests.post(
                f"{self.endpoint}/translate?api-version=3.0&from=en&to=hi",
                headers=self.headers,
                json=payload
            )
            if response.status_code == 200:
                return response.json()[0]["translations"][0]["text"]
            return text
        except Exception as e:
            print(f"Translation error: {e}")
            return text

class DocumentProcessor:
    def __init__(self):
        self.translator = AzureTranslator()

    def process_document(self, input_path: str, output_path: str) -> bool:
        try:
            doc = Document(input_path)
            total_items = len(doc.paragraphs) + sum(len(table.rows) * len(table.columns) for table in doc.tables)
            
            with tqdm(total=total_items, desc="Translating") as pbar:
                # Process tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                cell.text = self.translator.translate_text(cell.text)
                                pbar.update(1)

                # Process paragraphs (skip images)
                for para in doc.paragraphs:
                    if para.text.strip() and not para._element.xpath('.//w:drawing'):
                        para.text = self.translator.translate_text(para.text)
                    pbar.update(1)

            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Document processing error: {e}")
            return False

def main():
    processor = DocumentProcessor()
    input_file = "input.docx"
    output_file = "translated_hindi3333.docx"

    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found")
        return

    print("Starting translation...")
    if processor.process_document(input_file, output_file):
        print(f"Document translated successfully: {output_file}")
    else:
        print("Translation failed")

if __name__ == "__main__":
    main()