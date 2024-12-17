import os
import docx
from docx.shared import Pt
from langchain.chat_models import ChatOpenAI
from langchain.schema import HumanMessage, SystemMessage
from tqdm import tqdm

def translate_text_gpt4(content, llm):
    messages = [
        SystemMessage(content=(
            "You are a professional Hindi translator. Translate the following text into Hindi while retaining the meaning, "
            "context, and tone. However, do not translate any financial terms, terminology, or phrases. Keep all financial "
            "terms in English as is."
        )),
        HumanMessage(content=content)
    ]
    response = llm(messages)
    return response.content.strip()

def process_docx(input_file, output_file, llm):
    doc = docx.Document(input_file)
    translated_doc = docx.Document()
    total_items = len(doc.paragraphs) + sum(len(table.rows) for table in doc.tables)
    pbar = tqdm(total=total_items, desc="Translating")

    for element in doc.element.body:
        if element.tag.endswith('p'):
            para = docx.text.paragraph.Paragraph(element, doc)
            new_para = translated_doc.add_paragraph()
            for run in para.runs:
                text = run.text.strip()
                if text:
                    translated_text = translate_text_gpt4(text, llm)
                    new_run = new_para.add_run(translated_text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name
                else:
                    new_para.add_run()
            new_para.alignment = para.alignment
            pbar.update(1)
        elif element.tag.endswith('tbl'):
            table = docx.table.Table(element, doc)
            new_table = translated_doc.add_table(rows=0, cols=len(table.columns))
            for row in table.rows:
                new_row = new_table.add_row()
                for idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    translated_text = translate_text_gpt4(text, llm) if text else ''
                    new_row.cells[idx].text = translated_text
                pbar.update(1)
        else:
            translated_doc.element.body.append(element)
            pbar.update(1)

    pbar.close()
    translated_doc.save(output_file)
    print(f"Translated document saved to: {output_file}")

if __name__ == "__main__":
    os.environ["OPENAI_API_KEY"] = "your open ai key"  # Replace with your OpenAI API key
    llm = ChatOpenAI(model_name="gpt-4o", temperature=0)
    input_file = "input.docx"
    output_file = "output_hindi_custom_prompt.docx"
    if os.path.exists(input_file):
        process_docx(input_file, output_file, llm)
    else:
        print("Input file not found.")
