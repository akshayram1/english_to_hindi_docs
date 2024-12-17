import os
import time
from typing import Optional
from docx import Document
from docx.shared import Pt
from langchain.chat_models import ChatOpenAI
from langchain.schema import HumanMessage, SystemMessage
from tqdm import tqdm

class DocumentProcessor:
    def __init__(self, api_key: str, model_name: str = "gpt-4o"):
        self.llm = ChatOpenAI(
            api_key=api_key,
            model_name=model_name,
            temperature=0
        )
        self.max_retries = 3
        self.retry_delay = 2

    def translate_text(self, text: str) -> str:
        if not text.strip():
            return text
            
        messages = [
            SystemMessage(content="You are a professional Hindi translator. You are Finance Expert. You'll not translate Financial Terms and keep them as it is. Translate the following text into Hindi while retaining the meaning, context, and tone."),
            HumanMessage(content=text)
        ]
        
        for attempt in range(self.max_retries):
            try:
                response = self.llm(messages)
                return response.content.strip()
            except Exception as e:
                print(f"Translation attempt {attempt + 1} failed: {e}")
                if attempt < self.max_retries - 1:
                    time.sleep(self.retry_delay)
        return text

    def process_document(self, input_path: str, output_path: str) -> bool:
        try:
            doc = Document(input_path)
            total_items = len(doc.paragraphs) + sum(len(table.rows) for table in doc.tables)
            
            with tqdm(total=total_items, desc="Translating") as pbar:
                # Process tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                cell.text = self.translate_text(cell.text)
                        pbar.update(1)

                # Process paragraphs
                for para in doc.paragraphs:
                    if para.text.strip() and not para._element.xpath('.//w:drawing'):
                        for run in para.runs:
                            if run.text.strip():
                                original_format = {
                                    'bold': run.bold,
                                    'italic': run.italic,
                                    'font': run.font.name,
                                    'size': run.font.size
                                }
                                run.text = self.translate_text(run.text)
                                # Restore formatting
                                run.bold = original_format['bold']
                                run.italic = original_format['italic']
                                run.font.name = original_format['font']
                                run.font.size = original_format['size']
                    pbar.update(1)

            # Save document directly
            doc.save(output_path)
            print(f"Translation complete: {output_path}")
            return True

        except Exception as e:
            print(f"Document processing error: {e}")
            return False

def main():
    api_key = "your open ai key"
    processor = DocumentProcessor(api_key)
    
    input_file = "input.docx"
    output_file = "transllllllllll.docx"

    if os.path.exists(input_file):
        if processor.process_document(input_file, output_file):
            print("Document translated successfully")
        else:
            print("Translation failed")
    else:
        print("Input file not found")

if __name__ == "__main__":
    main()