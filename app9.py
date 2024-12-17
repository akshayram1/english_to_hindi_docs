import torch
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
from IndicTransTokenizer import IndicProcessor
from docx import Document
from tqdm import tqdm
import os

class IndicTranslator:
    def __init__(self):
        print("Loading AI4Bharat IndicTrans2 model... (this may take a few minutes)")
        self.model_name = "ai4bharat/indictrans2-en-indic-dist-200M"
        self.tokenizer = AutoTokenizer.from_pretrained(self.model_name, trust_remote_code=True)
        self.model = AutoModelForSeq2SeqLM.from_pretrained(self.model_name, trust_remote_code=True)
        self.ip = IndicProcessor(inference=True)
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.model.to(self.device)
        print(f"Model loaded successfully on {self.device}!")

    def translate_batch(self, texts):
        if not texts:
            return []
            
        try:
            batch = self.ip.preprocess_batch(texts, src_lang="eng_Latn", tgt_lang="hin_Deva")
            
            inputs = self.tokenizer(
                batch,
                truncation=True,
                padding="longest",
                return_tensors="pt",
                return_attention_mask=True
            ).to(self.device)

            with torch.no_grad():
                generated_tokens = self.model.generate(
                    **inputs,
                    use_cache=True,
                    min_length=0,
                    max_length=256,
                    num_beams=5,
                    num_return_sequences=1
                )

            with self.tokenizer.as_target_tokenizer():
                generated_tokens = self.tokenizer.batch_decode(
                    generated_tokens.detach().cpu().tolist(),
                    skip_special_tokens=True,
                    clean_up_tokenization_spaces=True
                )

            translations = self.ip.postprocess_batch(generated_tokens, lang="hin_Deva")
            return translations
        except Exception as e:
            print(f"Translation error: {e}")
            return texts

    def process_document(self, input_path: str, output_path: str) -> bool:
        try:
            doc = Document(input_path)
            
            # Count total items for progress bar
            total_items = len(doc.paragraphs) + sum(len(table.rows) * len(table.columns) for table in doc.tables)
            
            with tqdm(total=total_items, desc="Translating") as pbar:
                # Process tables
                for table in doc.tables:
                    texts = []
                    cells = []
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                texts.append(cell.text)
                                cells.append(cell)
                    
                    if texts:
                        translations = self.translate_batch(texts)
                        for cell, translation in zip(cells, translations):
                            cell.text = translation
                            pbar.update(1)

                # Process paragraphs
                texts = []
                paras = []
                for para in doc.paragraphs:
                    if para.text.strip() and not para._element.xpath('.//w:drawing'):
                        texts.append(para.text)
                        paras.append(para)
                
                if texts:
                    translations = self.translate_batch(texts)
                    for para, translation in zip(paras, translations):
                        para.text = translation
                        pbar.update(1)

            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Document processing error: {e}")
            return False

def main():
    translator = IndicTranslator()
    input_file = "input.docx"
    output_file = "translated_hindi_ai4bharat.docx"

    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found")
        return

    print("Starting translation...")
    if translator.process_document(input_file, output_file):
        print(f"Document translated successfully: {output_file}")
    else:
        print("Translation failed")

if __name__ == "__main__":
    main()