from googletrans import Translator
from docx import Document
from jina import Client
import time
import os
from typing import Optional

class DocumentTranslator:
    def __init__(self):
        self.translator = Translator()
        self.max_retries = 3
        self.retry_delay = 2

    def translate_text(self, text: str, dest='hi') -> str:
        if not text.strip():
            return text

        for attempt in range(self.max_retries):
            try:
                result = self.translator.translate(text, dest=dest)
                return result.text
            except Exception as e:
                print(f"Translation attempt {attempt + 1} failed: {e}")
                if attempt < self.max_retries - 1:
                    time.sleep(self.retry_delay)
        return text

    def process_document(self, input_path: str, output_path: str) -> bool:
        try:
            # Load document
            doc = Document(input_path)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            cell.text = self.translate_text(cell.text)

            # Process paragraphs (skip images)
            for para in doc.paragraphs:
                if para.text.strip() and not para._element.xpath('.//w:drawing'):
                    translated_text = self.translate_text(para.text)
                    para.text = translated_text

            # Save using Jina format
            doc.save(output_path)
            print(f"Translation complete: {output_path}")
            return True

        except Exception as e:
            print(f"Document processing error: {e}")
            return False

def main():
    translator = DocumentTranslator()
    input_file = "input3.docx"
    output_file = "translated_hindi.docx"

    if translator.process_document(input_file, output_file):
        print("Document translated successfully")
    else:
        print("Translation failed")

if __name__ == "__main__":
    main()