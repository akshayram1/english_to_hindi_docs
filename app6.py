from englisttohindi.englisttohindi import EngtoHindi
from docx import Document
import os

class LocalTranslator:
    def __init__(self):
        # Initialize without message
        self.translators = {}
        
    def translate_text(self, text: str) -> str:
        if not text.strip():
            return text
            
        try:
            # Create new translator instance for each text
            translator = EngtoHindi(text)
            return translator.convert
        except Exception as e:
            print(f"Translation error: {e}")
            return text

class DocumentTranslator:
    def __init__(self):
        self.translator = LocalTranslator()

    def process_document(self, input_path: str, output_path: str) -> bool:
        try:
            doc = Document(input_path)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            translated = self.translator.translate_text(cell.text)
                            cell.text = translated if translated else cell.text

            # Process paragraphs (skip images)
            for para in doc.paragraphs:
                if para.text.strip() and not para._element.xpath('.//w:drawing'):
                    translated = self.translator.translate_text(para.text)
                    para.text = translated if translated else para.text

            doc.save(output_path)
            print(f"Document saved to: {output_path}")
            return True
        except Exception as e:
            print(f"Document processing error: {e}")
            return False

def main():
    translator = DocumentTranslator()
    input_file = "input3.docx"
    output_file = "translated_hindi_input3.docx"

    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found")
        return

    if translator.process_document(input_file, output_file):
        print(f"Document translated successfully: {output_file}")
    else:
        print(f"Translation failed for {input_file}")

if __name__ == "__main__":
    main()