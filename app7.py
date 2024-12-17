import requests
import json
from docx import Document
import os
from typing import Optional
import time

class OTranslator:
    def __init__(self):
        self.api_key = ""
        self.base_url = 'https://otranslator.com/api/v1/translation'
        self.headers = {
            'Authorization': self.api_key,
            'Content-Type': 'application/json'
        }

    def create_translation_task(self, text: str) -> Optional[str]:
        try:
            data = {
                'texts': [text],
                'sourceLang': 'en',
                'targetLang': 'hi'
            }
            response = requests.post(
                f'{self.base_url}/create',
                headers=self.headers,
                json=data
            )
            if response.status_code == 200:
                return response.json().get('taskId')
            return None
        except Exception as e:
            print(f"Task creation error: {e}")
            return None

    def query_translation(self, task_id: str) -> Optional[str]:
        try:
            data = {'taskId': task_id}
            response = requests.post(
                f'{self.base_url}/queryTexts',
                headers=self.headers,
                json=data
            )
            if response.status_code == 200:
                translations = response.json().get('translations', {})
                return next(iter(translations.values()), None)
            return None
        except Exception as e:
            print(f"Query error: {e}")
            return None

    def translate_text(self, text: str) -> str:
        if not text.strip():
            return text

        task_id = self.create_translation_task(text)
        if not task_id:
            return text

        # Wait for translation to complete
        for _ in range(3):
            time.sleep(2)
            result = self.query_translation(task_id)
            if result:
                return result
        
        return text

class DocumentTranslator:
    def __init__(self):
        self.translator = OTranslator()

    def process_document(self, input_path: str, output_path: str) -> bool:
        try:
            doc = Document(input_path)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            cell.text = self.translator.translate_text(cell.text)

            # Process paragraphs (skip images)
            for para in doc.paragraphs:
                if para.text.strip() and not para._element.xpath('.//w:drawing'):
                    para.text = self.translator.translate_text(para.text)

            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Document processing error: {e}")
            return False

def main():
    translator = DocumentTranslator()
    input_file = "input3.docx"
    output_file = "translated_hindi_input333.docx"

    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found")
        return

    if translator.process_document(input_file, output_file):
        print(f"Document translated successfully: {output_file}")
    else:
        print(f"Translation failed for {input_file}")

if __name__ == "__main__":
    main()