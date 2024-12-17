import os
import time
import docx
from docx import Document
from langchain.chat_models import ChatOpenAI
from langchain.schema import HumanMessage, SystemMessage
from tqdm import tqdm

class DocumentTranslator:
    def __init__(self, llm):
        self.llm = llm
        self.max_retries = 3
        self.retry_delay = 2

    def translate_text_gpt4(self, content):
        messages = [
            SystemMessage(content=(
                "You are a professional Hindi translator. Translate the following text into Hindi while retaining the meaning, "
                "context, and tone. However, do not translate any financial terms, terminology, or phrases. Keep all financial "
                "terms in English as is."
            )),
            HumanMessage(content=content)
        ]
        response = self.llm(messages)
        return response.content.strip()

    def process_document(self, input_path: str, output_path: str) -> bool:
        try:
            # Load document
            doc = Document(input_path)
            translated_doc = Document()
            total_items = len(doc.paragraphs) + sum(len(table.rows) for table in doc.tables)
            pbar = tqdm(total=total_items, desc="Translating")

            # Process tables
            for table in doc.tables:
                new_table = translated_doc.add_table(rows=0, cols=len(table.columns))
                for row in table.rows:
                    new_row = new_table.add_row()
                    for idx, cell in enumerate(row.cells):
                        text = cell.text.strip()
                        translated_text = self.translate_text_gpt4(text) if text else ''
                        new_row.cells[idx].text = translated_text
                    pbar.update(1)

            # Process paragraphs (skip images)
            for para in doc.paragraphs:
                new_para = translated_doc.add_paragraph()
                for run in para.runs:
                    text = run.text.strip()
                    if text:
                        translated_text = self.translate_text_gpt4(text)
                        new_run = new_para.add_run(translated_text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.size = run.font.size
                        new_run.font.name = run.font.name
                    else:
                        new_para.add_run()
                new_para.alignment = para.alignment
                pbar.update(1)

            pbar.close()
            translated_doc.save(output_path)
            print(f"Translation complete: {output_path}")
            return True

        except Exception as e:
            print(f"Document processing error: {e}")
            return False

def main():
    os.environ["OPENAI_API_KEY"] = "your open ai key"  # Replace with your OpenAI API key
    llm = ChatOpenAI(model_name="gpt-4o", temperature=0)
    translator = DocumentTranslator(llm)
    input_file = "input.docx"
    output_file = "translatdi.docx"

    if os.path.exists(input_file):
        if translator.process_document(input_file, output_file):
            print("Document translated successfully")
        else:
            print("Translation failed")
    else:
        print("Input file not found")

if __name__ == "__main__":
    main()